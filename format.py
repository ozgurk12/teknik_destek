#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Okul Öncesi Etkinlik Rehberi Şablonu
Etkinlik planlama ve dokümantasyon için Python formatı
"""

class EtkinlikRehberi:
    def __init__(self):
        self.etkinlik_adi = ""
        self.alan_adi = ""  # Fen ve Ekoloji, Matematik, Sanat, vb.
        self.yas_grubu = ""  # örn: "36-48 ay"
        self.sure = ""  # örn: "30 dakika"
        self.uygulama_yeri = ""  # Sınıf içi, Bahçe, vb.
        
        self.amaclar = [
            "",
            "",
            "",
            ""
        ]
        
        self.materyaller = {
            "temel_malzemeler": "",
            "kategori_1": {
                "ana": [],
                "alternatif": []
            },
            "kategori_2": {
                "ana": [],
                "alternatif": []
            },
            "kategori_3": {
                "ana": [],
                "alternatif": []
            },
            "kategori_4": {
                "ana": [],
                "alternatif": []
            },
            "kategori_5": {
                "ana": [],
                "alternatif": []
            }
        }
        
    def giris_asamasi(self):
        """Etkinliğin giriş aşaması"""
        giris = {
            "sure": "",
            "adimlar": [
                "",
                "",
                "",
                "",
                "",
                ""
            ]
        }
        return giris
    
    def gelisme_asamasi(self):
        """Etkinliğin ana gelişme aşaması"""
        gelisme = {
            "sure": "",
            "aktiviteler": {
                "aktivite_1": "",
                "aktivite_2": "",
                "aktivite_3": "",
                "aktivite_4": "",
                "aktivite_5": ""
            }
        }
        return gelisme
    
    def yansima_cemberi(self):
        """Yansıma ve değerlendirme aşaması"""
        yansima = {
            "sure": "",
            "duzenleme": "",
            "sorular": [
                "",
                "",
                "",
                ""
            ],
            "ogretmen_ozet_ornekleri": [
                "",
                ""
            ],
            "kapaniş_sorusu": "",
            "final": ""
        }
        return yansima
    
    def sonuc_asamasi(self):
        """Etkinliğin sonuç/kapanış aşaması"""
        sonuc = {
            "sure": "",
            "aktiviteler": [
                "",
                ""
            ]
        }
        return sonuc
    
    def uyarlama_tablosu(self):
        """Özel gereksinimli çocuklar için uyarlama tablosu"""
        uyarlamalar = {
            "Görme Yetersizliği": "",
            "İşitme Yetersizliği": "",
            "Fiziksel/Motor Sınırlılık": "",
            "Dikkat Eksikliği/Hiperaktivite": "",
            "Dil ve Konuşma Güçlüğü": ""
        }
        return uyarlamalar
    
    def farklilastirma_ve_kapsayicilik(self):
        """Farklılaştırma ve kapsayıcılık stratejileri"""
        stratejiler = {
            "Öğrenme Hızı Farkı": {
                "ileri_duzey": "",
                "ek_destek": ""
            },
            "Dil Desteği": "",
            "Kültürel Kapsayıcılık": "",
            "Sosyal-Duygusal Farklılaştırma": "",
            "Çoklu Duyusal Yaklaşım": ""
        }
        return stratejiler
    
    def degerlendirme(self):
        """Etkinlik değerlendirme kriterleri"""
        degerlendirme = {
            "Program Tarafından": [
                "",
                "",
                ""
            ],
            "Amaçlanan Beceriler Tarafından": [
                "",
                "",
                ""
            ],
            "Öğrenciler Tarafından": [
                "",
                "",
                ""
            ]
        }
        return degerlendirme
    
    def etkinlik_dokumani_olustur(self):
        """Etkinlik dokümanını formatla ve oluştur"""
        dokuman = {
            "baslik": {
                "etkinlik_adi": self.etkinlik_adi,
                "alan_adi": self.alan_adi,
                "yas_grubu": self.yas_grubu,
                "sure": self.sure,
                "uygulama_yeri": self.uygulama_yeri
            },
            "amaclar": self.amaclar,
            "materyaller": self.materyaller,
            "uygulama_sureci": {
                "giris": self.giris_asamasi(),
                "gelisme": self.gelisme_asamasi(),
                "yansima": self.yansima_cemberi(),
                "sonuc": self.sonuc_asamasi()
            },
            "uyarlamalar": self.uyarlama_tablosu(),
            "farklilastirma": self.farklilastirma_ve_kapsayicilik(),
            "degerlendirme": self.degerlendirme()
        }
        return dokuman
    
    def yazdir_dokuman(self):
        """Etkinlik dokümanını konsola yazdır"""
        print("="*60)
        print(f"ETKİNLİK ADI: {self.etkinlik_adi}")
        print("="*60)
        print(f"Alan: {self.alan_adi}")
        print(f"Yaş Grubu: {self.yas_grubu}")
        print(f"Süre: {self.sure}")
        print(f"Yer: {self.uygulama_yeri}")
        
        print("\n" + "AMAÇLAR".center(60))
        print("-"*60)
        for i, amac in enumerate(self.amaclar, 1):
            if amac:
                print(f"{i}. {amac}")
        
        print("\n" + "MATERYALLER".center(60))
        print("-"*60)
        for kategori, icerik in self.materyaller.items():
            if isinstance(icerik, dict):
                print(f"\n{kategori.upper().replace('_', ' ')}:")
                if icerik['ana']:
                    print(f"  Ana: {', '.join(icerik['ana'])}")
                if icerik['alternatif']:
                    print(f"  Alternatif: {', '.join(icerik['alternatif'])}")
            else:
                if icerik:
                    print(f"{kategori}: {icerik}")
        
        print("\n" + "UYGULAMA SÜRECİ".center(60))
        print("="*60)
        
        # Giriş
        giris = self.giris_asamasi()
        print(f"\n1. GİRİŞ ({giris['sure']})")
        print("-"*40)
        for adim in giris['adimlar']:
            if adim:
                print(f"• {adim}")
        
        # Gelişme
        gelisme = self.gelisme_asamasi()
        print(f"\n2. GELİŞME ({gelisme['sure']})")
        print("-"*40)
        for aktivite, aciklama in gelisme['aktiviteler'].items():
            if aciklama:
                print(f"• {aktivite.upper().replace('_', ' ')}: {aciklama}")
        
        # Yansıma
        yansima = self.yansima_cemberi()
        print(f"\n3. YANSIMA ÇEMBERİ ({yansima['sure']})")
        print("-"*40)
        if yansima['duzenleme']:
            print(f"Düzenleme: {yansima['duzenleme']}")
        print("\nYönlendirici Sorular:")
        for soru in yansima['sorular']:
            if soru:
                print(f"• {soru}")
        if yansima['ogretmen_ozet_ornekleri'][0]:
            print("\nÖğretmen Özet Örnekleri:")
            for ornek in yansima['ogretmen_ozet_ornekleri']:
                if ornek:
                    print(f"• {ornek}")
        if yansima['kapaniş_sorusu']:
            print(f"\nKapanış Sorusu: {yansima['kapaniş_sorusu']}")
        if yansima['final']:
            print(f"Final: {yansima['final']}")
        
        # Sonuç
        sonuc = self.sonuc_asamasi()
        print(f"\n4. SONUÇ ({sonuc['sure']})")
        print("-"*40)
        for aktivite in sonuc['aktiviteler']:
            if aktivite:
                print(f"• {aktivite}")
        
        print("\n" + "UYARLAMA TABLOSU".center(60))
        print("="*60)
        uyarlamalar = self.uyarlama_tablosu()
        for gereksinim, uygulama in uyarlamalar.items():
            if uygulama:
                print(f"\n{gereksinim}:")
                print(f"  → {uygulama}")
        
        print("\n" + "FARKLILAŞTIRMA VE KAPSAYICILIK".center(60))
        print("="*60)
        stratejiler = self.farklilastirma_ve_kapsayicilik()
        for baslik, icerik in stratejiler.items():
            if isinstance(icerik, dict):
                print(f"\n{baslik}:")
                for alt_baslik, aciklama in icerik.items():
                    if aciklama:
                        print(f"  • {alt_baslik}: {aciklama}")
            else:
                if icerik:
                    print(f"\n{baslik}:")
                    print(f"  → {icerik}")
        
        print("\n" + "DEĞERLENDİRME".center(60))
        print("="*60)
        degerlendirme = self.degerlendirme()
        for perspektif, sorular in degerlendirme.items():
            print(f"\n{perspektif}:")
            for soru in sorular:
                if soru:
                    print(f"  • {soru}")
    
    def kaydet_json(self, dosya_adi="etkinlik.json"):
        """Etkinliği JSON formatında kaydet"""
        import json
        dokuman = self.etkinlik_dokumani_olustur()
        with open(dosya_adi, 'w', encoding='utf-8') as f:
            json.dump(dokuman, f, ensure_ascii=False, indent=2)
        print(f"Etkinlik {dosya_adi} dosyasına kaydedildi.")
    
    def kaydet_docx(self, dosya_adi="etkinlik.docx"):
        """Etkinliği DOCX formatında kaydet (python-docx gerektirir)"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            doc = Document()
            
            # Başlık
            baslik = doc.add_heading(f'Etkinlik Adı: {self.etkinlik_adi}', 0)
            baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Temel bilgiler
            doc.add_heading('Temel Bilgiler', 1)
            doc.add_paragraph(f'**Alan Adı:** {self.alan_adi}')
            doc.add_paragraph(f'**Yaş Grubu:** {self.yas_grubu}')
            doc.add_paragraph(f'**Süre:** {self.sure}')
            doc.add_paragraph(f'**Uygulama Yeri:** {self.uygulama_yeri}')
            
            # Amaçlar
            doc.add_heading('Etkinliğin Amacı', 1)
            for amac in self.amaclar:
                if amac:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(amac)
            
            # Materyaller
            doc.add_heading('Materyaller', 1)
            for kategori, icerik in self.materyaller.items():
                if isinstance(icerik, dict):
                    p = doc.add_paragraph()
                    p.add_run(f'{kategori.upper().replace("_", " ")}:').bold = True
                    if icerik['ana']:
                        doc.add_paragraph(f'• Ana: {", ".join(icerik["ana"])}')
                    if icerik['alternatif']:
                        doc.add_paragraph(f'• Alternatif: {", ".join(icerik["alternatif"])}')
                else:
                    if icerik:
                        p = doc.add_paragraph()
                        p.add_run(f'{kategori}: ').bold = True
                        p.add_run(icerik)
            
            # Uygulama Süreci
            doc.add_heading('Uygulama Süreci', 1)
            
            # Giriş
            giris = self.giris_asamasi()
            doc.add_heading(f'Giriş ({giris["sure"]})', 2)
            for adim in giris['adimlar']:
                if adim:
                    doc.add_paragraph(f'• {adim}')
            
            # Gelişme
            gelisme = self.gelisme_asamasi()
            doc.add_heading(f'Gelişme ({gelisme["sure"]})', 2)
            for aktivite, aciklama in gelisme['aktiviteler'].items():
                if aciklama:
                    p = doc.add_paragraph()
                    p.add_run(f'• {aktivite.upper().replace("_", " ")}: ').bold = True
                    p.add_run(aciklama)
            
            # Yansıma Çemberi
            yansima = self.yansima_cemberi()
            doc.add_heading(f'Yansıma Çemberi ({yansima["sure"]})', 2)
            if yansima['duzenleme']:
                doc.add_paragraph(f'Düzenleme: {yansima["duzenleme"]}')
            
            p = doc.add_paragraph()
            p.add_run('Yönlendirici Sorular:').bold = True
            for soru in yansima['sorular']:
                if soru:
                    doc.add_paragraph(f'• {soru}')
            
            if yansima['ogretmen_ozet_ornekleri'][0]:
                p = doc.add_paragraph()
                p.add_run('Öğretmen Özet Örnekleri:').bold = True
                for ornek in yansima['ogretmen_ozet_ornekleri']:
                    if ornek:
                        doc.add_paragraph(f'• {ornek}')
            
            if yansima['kapaniş_sorusu']:
                p = doc.add_paragraph()
                p.add_run('Kapanış Sorusu: ').bold = True
                p.add_run(yansima['kapaniş_sorusu'])
            
            if yansima['final']:
                p = doc.add_paragraph()
                p.add_run('Final: ').bold = True
                p.add_run(yansima['final'])
            
            # Sonuç
            sonuc = self.sonuc_asamasi()
            doc.add_heading(f'Sonuç ({sonuc["sure"]})', 2)
            for aktivite in sonuc['aktiviteler']:
                if aktivite:
                    doc.add_paragraph(f'• {aktivite}')
            
            # Uyarlama Tablosu
            doc.add_page_break()
            doc.add_heading('UYARLAMA', 1)
            
            # Tablo oluştur
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Başlık satırı
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Gereksinim Türü'
            hdr_cells[1].text = 'Uygulama Ayrıntısı'
            
            # Başlık satırını kalın yap
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Uyarlama verilerini ekle
            uyarlamalar = self.uyarlama_tablosu()
            for gereksinim, uygulama in uyarlamalar.items():
                if uygulama:
                    row_cells = table.add_row().cells
                    row_cells[0].text = gereksinim
                    row_cells[1].text = uygulama
            
            # Farklılaştırma ve Kapsayıcılık
            doc.add_heading('FARKLILAŞTIRMA VE KAPSAYICILIK', 1)
            stratejiler = self.farklilastirma_ve_kapsayicilik()
            
            for baslik, icerik in stratejiler.items():
                p = doc.add_paragraph()
                p.add_run(f'{baslik}:').bold = True
                
                if isinstance(icerik, dict):
                    for alt_baslik, aciklama in icerik.items():
                        if aciklama:
                            doc.add_paragraph(f'• {alt_baslik}: {aciklama}')
                else:
                    if icerik:
                        doc.add_paragraph(f'• {icerik}')
            
            # Değerlendirme
            doc.add_heading('DEĞERLENDİRME', 1)
            degerlendirme = self.degerlendirme()
            
            for perspektif, sorular in degerlendirme.items():
                p = doc.add_paragraph()
                p.add_run(f'{perspektif}:').bold = True
                for soru in sorular:
                    if soru:
                        doc.add_paragraph(f'• {soru}')
            
            # Belgeyi kaydet
            doc.save(dosya_adi)
            print(f"Etkinlik {dosya_adi} dosyasına kaydedildi.")
            
        except ImportError:
            print("python-docx modülü yüklü değil. pip install python-docx ile yükleyebilirsiniz.")
    
    def kaydet_html(self, dosya_adi="etkinlik.html"):
        """Etkinliği HTML formatında kaydet"""
        html_content = f"""
<!DOCTYPE html>
<html lang="tr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self.etkinlik_adi} - Etkinlik Rehberi</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f4f4f4;
        }}
        h1 {{
            color: #333;
            text-align: center;
            border-bottom: 3px solid #4CAF50;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #4CAF50;
            margin-top: 30px;
            border-bottom: 1px solid #ddd;
            padding-bottom: 5px;
        }}
        h3 {{
            color: #666;
            margin-top: 20px;
        }}
        .info-box {{
            background: white;
            padding: 15px;
            border-radius: 5px;
            margin: 10px 0;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        .info-box strong {{
            color: #4CAF50;
        }}
        ul {{
            background: white;
            padding: 20px;
            border-radius: 5px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            background: white;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        th, td {{
            padding: 12px;
            text-align: left;
            border: 1px solid #ddd;
        }}
        th {{
            background-color: #4CAF50;
            color: white;
        }}
        tr:nth-child(even) {{
            background-color: #f9f9f9;
        }}
    </style>
</head>
<body>
    <h1>Etkinlik Adı: {self.etkinlik_adi}</h1>
    
    <div class="info-box">
        <p><strong>Alan Adı:</strong> {self.alan_adi}</p>
        <p><strong>Yaş Grubu:</strong> {self.yas_grubu}</p>
        <p><strong>Süre:</strong> {self.sure}</p>
        <p><strong>Uygulama Yeri:</strong> {self.uygulama_yeri}</p>
    </div>
    
    <h2>Etkinliğin Amacı</h2>
    <ul>
        {''.join([f'<li>{amac}</li>' for amac in self.amaclar if amac])}
    </ul>
    
    <h2>Materyaller</h2>
    <div class="info-box">
"""
        
        for kategori, icerik in self.materyaller.items():
            if isinstance(icerik, dict):
                if icerik['ana'] or icerik['alternatif']:
                    html_content += f"<h3>{kategori.upper().replace('_', ' ')}</h3>"
                    if icerik['ana']:
                        html_content += f"<p><strong>Ana:</strong> {', '.join(icerik['ana'])}</p>"
                    if icerik['alternatif']:
                        html_content += f"<p><strong>Alternatif:</strong> {', '.join(icerik['alternatif'])}</p>"
            else:
                if icerik:
                    html_content += f"<p><strong>{kategori}:</strong> {icerik}</p>"
        
        html_content += """
    </div>
    
    <h2>Uygulama Süreci</h2>
"""
        
        # Giriş
        giris = self.giris_asamasi()
        html_content += f"""
    <h3>Giriş ({giris['sure']})</h3>
    <ul>
        {''.join([f'<li>{adim}</li>' for adim in giris['adimlar'] if adim])}
    </ul>
"""
        
        # Gelişme
        gelisme = self.gelisme_asamasi()
        html_content += f"""
    <h3>Gelişme ({gelisme['sure']})</h3>
    <ul>
"""
        for aktivite, aciklama in gelisme['aktiviteler'].items():
            if aciklama:
                html_content += f"<li><strong>{aktivite.upper().replace('_', ' ')}:</strong> {aciklama}</li>"
        html_content += "</ul>"
        
        # Yansıma
        yansima = self.yansima_cemberi()
        html_content += f"""
    <h3>Yansıma Çemberi ({yansima['sure']})</h3>
    <div class="info-box">
"""
        if yansima['duzenleme']:
            html_content += f"<p><strong>Düzenleme:</strong> {yansima['duzenleme']}</p>"
        
        html_content += "<p><strong>Yönlendirici Sorular:</strong></p><ul>"
        for soru in yansima['sorular']:
            if soru:
                html_content += f"<li>{soru}</li>"
        html_content += "</ul>"
        
        if yansima['kapaniş_sorusu']:
            html_content += f"<p><strong>Kapanış Sorusu:</strong> {yansima['kapaniş_sorusu']}</p>"
        if yansima['final']:
            html_content += f"<p><strong>Final:</strong> {yansima['final']}</p>"
        html_content += "</div>"
        
        # Sonuç
        sonuc = self.sonuc_asamasi()
        html_content += f"""
    <h3>Sonuç ({sonuc['sure']})</h3>
    <ul>
        {''.join([f'<li>{aktivite}</li>' for aktivite in sonuc['aktiviteler'] if aktivite])}
    </ul>
"""
        
        # Uyarlama Tablosu
        html_content += """
    <h2>Uyarlama Tablosu</h2>
    <table>
        <tr>
            <th>Gereksinim Türü</th>
            <th>Uygulama Ayrıntısı</th>
        </tr>
"""
        uyarlamalar = self.uyarlama_tablosu()
        for gereksinim, uygulama in uyarlamalar.items():
            if uygulama:
                html_content += f"""
        <tr>
            <td>{gereksinim}</td>
            <td>{uygulama}</td>
        </tr>
"""
        html_content += "</table>"
        
        # Farklılaştırma
        html_content += """
    <h2>Farklılaştırma ve Kapsayıcılık</h2>
    <div class="info-box">
"""
        stratejiler = self.farklilastirma_ve_kapsayicilik()
        for baslik, icerik in stratejiler.items():
            if isinstance(icerik, dict):
                html_content += f"<h3>{baslik}</h3><ul>"
                for alt_baslik, aciklama in icerik.items():
                    if aciklama:
                        html_content += f"<li><strong>{alt_baslik}:</strong> {aciklama}</li>"
                html_content += "</ul>"
            else:
                if icerik:
                    html_content += f"<p><strong>{baslik}:</strong> {icerik}</p>"
        html_content += "</div>"
        
        # Değerlendirme
        html_content += """
    <h2>Değerlendirme</h2>
    <div class="info-box">
"""
        degerlendirme = self.degerlendirme()
        for perspektif, sorular in degerlendirme.items():
            html_content += f"<h3>{perspektif}</h3><ul>"
            for soru in sorular:
                if soru:
                    html_content += f"<li>{soru}</li>"
            html_content += "</ul>"
        html_content += """
    </div>
</body>
</html>
"""
        
        with open(dosya_adi, 'w', encoding='utf-8') as f:
            f.write(html_content)
        print(f"Etkinlik {dosya_adi} dosyasına kaydedildi.")


# Kullanım örneği
if __name__ == "__main__":
    # Yeni etkinlik oluştur
    etkinlik = EtkinlikRehberi()
    
    # Temel bilgileri doldur
    etkinlik.etkinlik_adi = "Etkinlik Adı Buraya"
    etkinlik.alan_adi = "Alan Adı"
    etkinlik.yas_grubu = "Yaş Grubu"
    etkinlik.sure = "Süre"
    etkinlik.uygulama_yeri = "Uygulama Yeri"
    
    # Amaçları doldur
    etkinlik.amaclar = [
        "Amaç 1",
        "Amaç 2",
        "Amaç 3",
        "Amaç 4"
    ]
    
    # Dokümanı yazdır
    etkinlik.yazdir_dokuman()
    
    # İsteğe bağlı: JSON olarak kaydet
    # etkinlik.kaydet_json()
    
    # İsteğe bağlı: DOCX olarak kaydet
    # etkinlik.kaydet_docx()
    
    # İsteğe bağlı: HTML olarak kaydet
    # etkinlik.kaydet_html()