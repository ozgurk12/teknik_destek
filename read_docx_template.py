from docx import Document
import json

def read_activity_template():
    """Örnek etkinlik DOCX dosyasını okur ve yapısını analiz eder"""
    
    doc_path = "Örnek Detaylı Etkinlik Rehberi.docx"
    doc = Document(doc_path)
    
    print("📄 DOCX DOSYASI YAPISI ANALİZİ\n")
    print("="*50)
    
    structure = {
        "paragraphs": [],
        "tables": [],
        "sections": []
    }
    
    # Paragrafları oku
    print("\n📝 PARAGRAFLAR:")
    print("-"*30)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"{i+1}. {para.text[:100]}...")
            structure["paragraphs"].append({
                "index": i,
                "text": para.text,
                "style": para.style.name if para.style else None,
                "alignment": str(para.alignment) if para.alignment else None
            })
    
    # Tabloları oku
    print(f"\n📊 TABLOLAR: {len(doc.tables)} adet")
    print("-"*30)
    for i, table in enumerate(doc.tables):
        print(f"\nTablo {i+1}: {len(table.rows)} satır x {len(table.columns)} sütun")
        table_data = []
        for j, row in enumerate(table.rows[:3]):  # İlk 3 satırı göster
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            print(f"  Satır {j+1}: {row_data}")
            table_data.append(row_data)
        
        structure["tables"].append({
            "index": i,
            "rows": len(table.rows),
            "columns": len(table.columns),
            "sample_data": table_data
        })
    
    # Yapıyı JSON olarak kaydet
    with open("docx_structure.json", "w", encoding="utf-8") as f:
        json.dump(structure, f, ensure_ascii=False, indent=2)
    
    print("\n✅ Yapı analizi 'docx_structure.json' dosyasına kaydedildi")
    
    return structure

if __name__ == "__main__":
    read_activity_template()