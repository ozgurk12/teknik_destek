# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from io import BytesIO
import json
from typing import Dict, Optional
import logging
from pathlib import Path
import os

logger = logging.getLogger(__name__)

class DocxExportService:
    """Service for exporting activities to DOCX format with background"""

    def __init__(self):
        # Color scheme from template
        self.primary_color = RGBColor(114, 191, 190)  # Turquoise/Teal
        self.text_color = RGBColor(44, 62, 80)  # Dark blue-gray
        self.background_image_path = '/Users/ozgurk12/Desktop/PROJELER/EduPage Kids/Etkinlik ve Günlük Plan/Varlık 2@3x.png'

    def add_watermark_to_page(self, doc):
        """Add watermark image directly to document body as background"""
        try:
            if not os.path.exists(self.background_image_path):
                logger.warning(f"Background image not found")
                return

            # Add the image as the first paragraph in document
            first_para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
            first_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = first_para.add_run()

            # Add the picture
            picture = run.add_picture(self.background_image_path, width=Cm(19))

            # Access the drawing element
            drawing = run._element.find('.//w:drawing', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if drawing is None:
                return

            # Get the inline element
            inline = drawing.find('.//wp:inline', namespaces={'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
            if inline is None:
                return

            # Create anchor element to float the image
            anchor = OxmlElement('wp:anchor')

            # Set anchor attributes for behind text
            anchor.set('distT', '0')
            anchor.set('distB', '0')
            anchor.set('distL', '0')
            anchor.set('distR', '0')
            anchor.set('simplePos', '0')
            anchor.set('relativeHeight', '0')
            anchor.set('behindDoc', '1')  # Behind text
            anchor.set('locked', '0')
            anchor.set('layoutInCell', '0')
            anchor.set('allowOverlap', '1')

            # Add positioning elements
            simplePos = OxmlElement('wp:simplePos')
            simplePos.set('x', '0')
            simplePos.set('y', '0')
            anchor.append(simplePos)

            # Horizontal position
            positionH = OxmlElement('wp:positionH')
            positionH.set('relativeFrom', 'page')
            posOffset = OxmlElement('wp:posOffset')
            posOffset.text = '0'
            positionH.append(posOffset)
            anchor.append(positionH)

            # Vertical position
            positionV = OxmlElement('wp:positionV')
            positionV.set('relativeFrom', 'page')
            posOffsetV = OxmlElement('wp:posOffset')
            posOffsetV.text = '0'
            positionV.append(posOffsetV)
            anchor.append(positionV)

            # Copy extent from inline
            for child in inline:
                if 'extent' in child.tag:
                    anchor.append(child)
                elif 'effectExtent' in child.tag:
                    anchor.append(child)
                elif 'docPr' in child.tag:
                    anchor.append(child)
                elif 'cNvGraphicFramePr' in child.tag:
                    anchor.append(child)
                elif 'graphic' in child.tag:
                    anchor.append(child)

            # Add wrap none
            wrapNone = OxmlElement('wp:wrapNone')
            anchor.append(wrapNone)

            # Replace inline with anchor
            drawing.remove(inline)
            drawing.append(anchor)

        except Exception as e:
            logger.error(f"Error adding watermark: {e}")

    def parse_text_to_list(self, text: str) -> list:
        """Parse bullet-pointed text into a list"""
        if not text:
            return []

        lines = text.split('\n')
        result = []
        for line in lines:
            line = line.strip()
            if line.startswith('•'):
                result.append(line[1:].strip())
            elif line:
                result.append(line)
        return result

    def export_activity_to_docx(self, activity_data: Dict) -> BytesIO:
        """Export activity data to DOCX format with background"""
        try:
            # Use activity_data directly
            json_data = activity_data

            # Create document
            doc = Document()

            # Set page margins
            for section in doc.sections:
                section.top_margin = Cm(2.5)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)
                # A4 size
                section.page_height = Cm(29.7)
                section.page_width = Cm(21)

            # Add first paragraph for watermark
            doc.add_paragraph()

            # Add watermark as background
            self.add_watermark_to_page(doc)

            # Set default font
            doc.styles['Normal'].font.name = 'Arial'
            doc.styles['Normal'].font.color.rgb = self.text_color

            # Title - add spacing first
            doc.add_paragraph()
            doc.add_paragraph()

            title_para = doc.add_paragraph()
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title_para.add_run(json_data.get('etkinlik_adi', 'ETKİNLİK').upper())
            title_run.font.name = 'Arial'
            title_run.font.size = Pt(24)
            title_run.font.bold = True
            title_run.font.color.rgb = self.primary_color

            doc.add_paragraph()

            # Basic Information Table
            para = doc.add_paragraph()
            run = para.add_run('Temel Bilgiler')
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = self.text_color

            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Table Grid'

            info_data = [
                ('Alan Adı:', json_data.get('alan_adi', '')),
                ('Yaş Grubu:', json_data.get('yas_grubu', '')),
                ('Süre:', f"{json_data.get('sure', '30')} dakika"),
                ('Uygulama Yeri:', json_data.get('uygulama_yeri', ''))
            ]

            for i, (label, value) in enumerate(info_data):
                row = info_table.rows[i]
                # Label
                cell_para = row.cells[0].paragraphs[0]
                cell_run = cell_para.add_run(label)
                cell_run.font.name = 'Arial'
                cell_run.font.size = Pt(11)
                cell_run.font.bold = True
                cell_run.font.color.rgb = self.text_color
                # Value
                cell_para = row.cells[1].paragraphs[0]
                cell_run = cell_para.add_run(value)
                cell_run.font.name = 'Arial'
                cell_run.font.size = Pt(11)
                cell_run.font.color.rgb = self.text_color

            # Objectives
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run('Etkinliğin Amacı:')
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = self.text_color

            amaclar = self.parse_text_to_list(json_data.get('etkinlik_amaci', ''))
            for amac in amaclar:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(amac)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = self.text_color

            # Materials
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run('Materyaller')
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = self.text_color

            materyaller = self.parse_text_to_list(json_data.get('materyaller', ''))
            for item in materyaller:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(item)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = self.text_color

            # Implementation Process
            doc.add_page_break()
            # Add watermark to new page
            doc.add_paragraph()
            self.add_watermark_to_page(doc)

            doc.add_paragraph()
            heading = doc.add_heading('UYGULAMA SÜRECİ', level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in heading.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(20)
                run.font.color.rgb = self.primary_color

            surec_text = json_data.get('uygulama_sureci', '')
            if surec_text:
                sections = surec_text.split('\n\n')
                for section in sections:
                    if section.strip():
                        lines = section.strip().split('\n')
                        if lines:
                            # Section header
                            header = lines[0]
                            para = doc.add_paragraph()
                            run = para.add_run(header)
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                            run.font.bold = True
                            run.font.color.rgb = self.text_color

                            # Content
                            for line in lines[1:]:
                                line = line.strip()
                                if line.startswith('•'):
                                    p = doc.add_paragraph(style='List Bullet')
                                    run = p.add_run(line[1:].strip())
                                elif line:
                                    p = doc.add_paragraph()
                                    run = p.add_run(line)
                                if line:
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = self.text_color

            # Adaptations
            if json_data.get('uyarlama'):
                doc.add_page_break()
                doc.add_paragraph()
                self.add_watermark_to_page(doc)

                heading = doc.add_heading('UYARLAMA VE FARKLILAŞTIRMA', level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in heading.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(20)
                    run.font.color.rgb = self.primary_color

                uyarlama_text = json_data.get('uyarlama', '')
                if uyarlama_text:
                    sections = uyarlama_text.split('\n\n')
                    for section in sections:
                        if section.strip():
                            lines = section.strip().split('\n')
                            for line in lines:
                                line = line.strip()
                                if line.endswith(':'):
                                    p = doc.add_paragraph()
                                    run = p.add_run(line)
                                    run.font.bold = True
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = self.text_color
                                elif line.startswith('•'):
                                    p = doc.add_paragraph(style='List Bullet')
                                    run = p.add_run(line[1:].strip())
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = self.text_color
                                elif line:
                                    p = doc.add_paragraph()
                                    run = p.add_run(line)
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = self.text_color

                # Farklılaştırma
                fark_text = json_data.get('farklilastirma_kapsayicilik', '')
                if fark_text:
                    doc.add_paragraph()
                    para = doc.add_paragraph()
                    run = para.add_run('Farklılaştırma ve Kapsayıcılık')
                    run.font.name = 'Arial'
                    run.font.size = Pt(14)
                    run.font.bold = True
                    run.font.color.rgb = self.text_color

                    sections = fark_text.split('\n\n')
                    for section in sections:
                        if section.strip():
                            lines = section.strip().split('\n')
                            for line in lines:
                                line = line.strip()
                                if line.endswith(':'):
                                    p = doc.add_paragraph()
                                    run = p.add_run(line)
                                    run.font.bold = True
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = self.text_color
                                elif line.startswith('•'):
                                    p = doc.add_paragraph(style='List Bullet')
                                    run = p.add_run(line[1:].strip())
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = self.text_color
                                elif line:
                                    p = doc.add_paragraph()
                                    run = p.add_run(line)
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = self.text_color

            # Evaluation
            if any([json_data.get('degerlendirme_program'),
                    json_data.get('degerlendirme_beceriler'),
                    json_data.get('degerlendirme_ogrenciler')]):
                doc.add_page_break()
                doc.add_paragraph()
                self.add_watermark_to_page(doc)

                heading = doc.add_heading('DEĞERLENDİRME', level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in heading.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(20)
                    run.font.color.rgb = self.primary_color

                eval_fields = [
                    ('degerlendirme_program', 'Program Değerlendirmesi'),
                    ('degerlendirme_beceriler', 'Beceri Değerlendirmesi'),
                    ('degerlendirme_ogrenciler', 'Öğrenci Değerlendirmesi')
                ]

                for field_key, field_title in eval_fields:
                    field_text = json_data.get(field_key, '')
                    if field_text:
                        para = doc.add_paragraph()
                        run = para.add_run(field_title + ':')
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
                        run.font.bold = True
                        run.font.color.rgb = self.text_color

                        lines = field_text.split('\n')
                        for line in lines:
                            line = line.strip()
                            if line.endswith(':'):
                                p = doc.add_paragraph()
                                run = p.add_run(line)
                                run.font.bold = True
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                                run.font.color.rgb = self.text_color
                            elif line.startswith('•'):
                                p = doc.add_paragraph(style='List Bullet')
                                run = p.add_run(line[1:].strip())
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                                run.font.color.rgb = self.text_color
                            elif line:
                                p = doc.add_paragraph()
                                run = p.add_run(line)
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                                run.font.color.rgb = self.text_color

                        doc.add_paragraph()

            # Save to BytesIO
            docx_file = BytesIO()
            doc.save(docx_file)
            docx_file.seek(0)

            return docx_file

        except Exception as e:
            logger.error(f"Error exporting to DOCX: {str(e)}")
            raise

# Singleton instance
docx_export_service = DocxExportService()