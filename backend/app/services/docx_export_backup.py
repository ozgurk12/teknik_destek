# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
from io import BytesIO
import json
from typing import Dict, Optional
import logging
from pathlib import Path
import os
import base64

logger = logging.getLogger(__name__)

class DocxExportService:
    """Service for exporting activities to DOCX format"""

    def __init__(self):
        # Color scheme from template
        self.primary_color = RGBColor(114, 191, 190)  # Turquoise/Teal
        self.text_color = RGBColor(44, 62, 80)  # Dark blue-gray
        self.header_bg_color = RGBColor(162, 216, 216)  # Light turquoise
        # Path to background image
        self.background_image_path = '/Users/ozgurk12/Desktop/PROJELER/EduPage Kids/Etkinlik ve Günlük Plan/Varlık 2@3x.png'

    def set_cell_background(self, cell, color_hex):
        """Set background color for a table cell"""
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), color_hex)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def add_watermark_to_section(self, section, doc):
        """Add image as watermark to section header"""
        try:
            if not os.path.exists(self.background_image_path):
                logger.warning(f"Background image not found at {self.background_image_path}")
                return

            # Get or create header
            header = section.header

            # Clear existing content
            for p in header.paragraphs:
                p.clear()

            # Add paragraph for watermark
            p = header.add_paragraph()
            r = p.add_run()

            # Add the background image
            r.add_picture(self.background_image_path, width=Cm(21))

            # Try to position it behind text using XML manipulation
            drawing = r._element.xpath('.//w:drawing')[0]

            # Create anchor element for floating image
            anchor = OxmlElement('wp:anchor')
            anchor.set('behindDoc', '1')  # Behind text
            anchor.set('locked', '0')
            anchor.set('layoutInCell', '1')
            anchor.set('allowOverlap', '1')

            # Copy drawing contents to anchor
            for child in drawing:
                anchor.append(child)

            # Replace drawing with anchor
            r._element.replace(drawing, anchor)

        except Exception as e:
            logger.warning(f"Could not add watermark: {e}")
            # Fallback: just add as header image
            try:
                header = section.header
                p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run()
                run.add_picture(self.background_image_path, width=Cm(19))
            except:
                pass

    def add_page_background(self, section):
        """Add background color to page section"""
        # Skip adding background color since we'll use the image
        pass

    def add_styled_header(self, doc, title, page_break=False):
        """Add a styled header"""
        if page_break:
            doc.add_page_break()
            # Add watermark to new section if created
            if len(doc.sections) > 1:
                self.add_watermark_to_section(doc.sections[-1], doc)

        # Add title as simple heading
        doc.add_paragraph()  # Spacing
        heading = doc.add_heading(title.upper(), level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in heading.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(20)
            run.font.color.rgb = self.primary_color
        doc.add_paragraph()  # Spacing after header

    def parse_text_to_list(self, text: str) -> list:
        """Parse bullet-pointed text into a list"""
        if not text:
            return []

        lines = text.split('\n')
        result = []
        for line in lines:
            line = line.strip()
            if line.startswith('•'):
                result.append(line[1:].strip())
            elif line:
                result.append(line)
        return result

    def export_activity_to_docx(self, activity_data: Dict) -> BytesIO:
        """Export activity data to DOCX format with proper UTF-8 encoding"""
        try:
            # Use activity_data directly as it's already flattened
            json_data = activity_data

            # Create a new Document
            doc = Document()

            # Configure page margins for better layout
            for section in doc.sections:
                section.top_margin = Cm(2.5)
                section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)

            # Add watermark to the section
            self.add_watermark_to_section(doc.sections[0], doc)

            # Set default font for better Turkish character support
            doc.styles['Normal'].font.name = 'Arial'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            doc.styles['Normal'].font.color.rgb = self.text_color
            
            # Add some spacing at the top
            doc.add_paragraph()
            doc.add_paragraph()

            # Activity title
            title_para = doc.add_paragraph()
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title_para.add_run(json_data.get('etkinlik_adi', 'ETKİNLİK').upper())
            title_run.font.name = 'Arial'
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title_run.font.color.rgb = self.primary_color

            doc.add_paragraph()  # Spacing
            
            # Basic Information
            para = doc.add_paragraph()
            run = para.add_run('Temel Bilgiler')
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = self.text_color
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ('Alan Adı:', json_data.get('alan_adi', '')),
                ('Yaş Grubu:', json_data.get('yas_grubu', '')),
                ('Süre:', f"{json_data.get('sure', '30')} dakika"),
                ('Uygulama Yeri:', json_data.get('uygulama_yeri', ''))
            ]
            
            for i, (label, value) in enumerate(info_data):
                row = info_table.rows[i]
                # Label cell
                cell_para = row.cells[0].paragraphs[0]
                cell_run = cell_para.add_run(label)
                cell_run.font.name = 'Arial'
                cell_run.font.size = Pt(11)
                cell_run.font.bold = True
                cell_run.font.color.rgb = self.text_color
                # Value cell
                cell_para = row.cells[1].paragraphs[0]
                cell_run = cell_para.add_run(value)
                cell_run.font.name = 'Arial'
                cell_run.font.size = Pt(11)
                cell_run.font.color.rgb = self.text_color
            
            # Objectives
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run('Etkinliğin Amacı:')
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = self.text_color

            # Parse etkinlik_amaci text into list
            amaclar = self.parse_text_to_list(json_data.get('etkinlik_amaci', ''))
            for amac in amaclar:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(amac)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = self.text_color
            
            # Materials
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run('Materyaller')
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = self.text_color

            # Parse materyaller text into list
            materyaller = self.parse_text_to_list(json_data.get('materyaller', ''))
            for item in materyaller:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(item)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = self.text_color
            
            # Implementation Process
            self.add_styled_header(doc, 'UYGULAMA SÜRECİ', page_break=True)

            # Parse uygulama_sureci text
            surec_text = json_data.get('uygulama_sureci', '')
            if surec_text:
                # Split by sections (GİRİŞ, GELİŞME, etc.)
                sections = surec_text.split('\n\n')
                for section in sections:
                    if section.strip():
                        lines = section.strip().split('\n')
                        if lines:
                            # First line is usually the section header
                            header = lines[0]
                            para = doc.add_paragraph()
                            run = para.add_run(header)
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                            run.font.bold = True
                            run.font.color.rgb = self.text_color

                            # Rest are content lines
                            for line in lines[1:]:
                                line = line.strip()
                                if line.startswith('•'):
                                    p = doc.add_paragraph(style='List Bullet')
                                    run = p.add_run(line[1:].strip())
                                elif line:
                                    p = doc.add_paragraph()
                                    run = p.add_run(line)
                                if line:
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = self.text_color

            # Legacy JSON format support
            surec = json_data.get('uygulama_sureci', {})
            if isinstance(surec, dict) and surec:
                # Giriş
                giris = surec.get('giris', {})
                if giris:
                    doc.add_heading(f"Giriş ({giris.get('sure', '5-7 dakika')})", 2)
                    for adim in giris.get('adimlar', []):
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(adim)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.color.rgb = self.text_color

                # Gelişme
                gelisme = surec.get('gelisme', {})
                if gelisme:
                    doc.add_heading(f"Gelişme ({gelisme.get('sure', '20 dakika')})", 2)
                    for key, value in gelisme.get('aktiviteler', {}).items():
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(value)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.color.rgb = self.text_color

                # Yansıma Çemberi
                yansima = surec.get('yansima_cemberi', {})
                if yansima:
                    doc.add_heading(f"Yansıma Çemberi ({yansima.get('sure', '5-7 dakika')})", 2)

                    if yansima.get('duzenleme'):
                        p = doc.add_paragraph()
                        p.add_run('Düzenleme: ').bold = True
                        p.add_run(yansima['duzenleme'])

                    if yansima.get('sorular'):
                        p = doc.add_paragraph()
                        p.add_run('Yönlendirici Sorular:').bold = True
                        for soru in yansima['sorular']:
                            p = doc.add_paragraph(style='List Bullet')
                            run = p.add_run(soru)
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                            run.font.color.rgb = self.text_color

                    if yansima.get('ogretmen_ozet_ornekleri'):
                        p = doc.add_paragraph()
                        p.add_run('Öğretmen Özet Örnekleri:').bold = True
                        for ornek in yansima['ogretmen_ozet_ornekleri']:
                            p = doc.add_paragraph(style='List Bullet')
                            run = p.add_run(ornek)
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                            run.font.color.rgb = self.text_color

                    if yansima.get('kapaniş_sorusu'):
                        p = doc.add_paragraph()
                        p.add_run('Kapanış Sorusu: ').bold = True
                        p.add_run(yansima['kapaniş_sorusu'])

                    if yansima.get('final'):
                        p = doc.add_paragraph()
                        p.add_run('Final: ').bold = True
                        p.add_run(yansima['final'])

                # Sonuç
                sonuc = surec.get('sonuc', {})
                if sonuc:
                    doc.add_heading(f"Sonuç ({sonuc.get('sure', '3 dakika')})", 2)
                    for aktivite in sonuc.get('aktiviteler', []):
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(aktivite)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.color.rgb = self.text_color
            
            # Adaptations
            self.add_styled_header(doc, 'UYARLAMA VE FARKLILAŞTIRMA', page_break=True)

            # Turkish field name mappings for uyarlama
            uyarlama_field_map = {
                'gorme_yetersizligi': 'Görme Yetersizliği',
                'isitme_yetersizligi': 'İşitme Yetersizliği',
                'fiziksel_motor_sinirlilik': 'Fiziksel Motor Sınırlılık',
                'dikkat_eksikligi_hiperaktivite': 'Dikkat Eksikliği ve Hiperaktivite',
                'otizm_spektrum_bozuklugu': 'Otizm Spektrum Bozukluğu',
                'dil_ve_konusma_guclugu': 'Dil ve Konuşma Güçlüğü',
                'zihinsel_yetersizlik': 'Zihinsel Yetersizlik',
                'ogrenme_guclugu': 'Öğrenme Güçlüğü'
            }

            # Parse uyarlama text
            uyarlama_text = json_data.get('uyarlama', '')
            if uyarlama_text and isinstance(uyarlama_text, str):
                para = doc.add_paragraph()
                run = para.add_run('Uyarlama')
                run.font.name = 'Arial'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = self.text_color

                # Split and format the text
                sections = uyarlama_text.split('\n\n')
                for section in sections:
                    if section.strip():
                        lines = section.strip().split('\n')
                        for line in lines:
                            line = line.strip()
                            if line.endswith(':'):
                                p = doc.add_paragraph()
                                run = p.add_run(line)
                                run.font.bold = True
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                                run.font.color.rgb = self.text_color
                            elif line.startswith('•'):
                                p = doc.add_paragraph(style='List Bullet')
                                run = p.add_run(line[1:].strip())
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                                run.font.color.rgb = self.text_color
                            elif line:
                                p = doc.add_paragraph()
                                run = p.add_run(line)
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                                run.font.color.rgb = self.text_color

            # Legacy JSON format
            uyarlama = json_data.get('uyarlama', {})
            if isinstance(uyarlama, dict) and uyarlama:
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Gereksinim Türü'
                hdr_cells[1].text = 'Uygulama Ayrıntısı'

                # Make header bold
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # Add data rows
                for key, value in uyarlama.items():
                    if value:
                        row_cells = table.add_row().cells
                        # If key already has Turkish characters, use it as is
                        if any(char in key for char in 'ÇĞÖŞÜçğöşüıİ'):
                            row_cells[0].text = key
                        else:
                            row_cells[0].text = uyarlama_field_map.get(key, key.replace('_', ' ').title())
                        row_cells[1].text = value
            
            # Differentiation and Inclusion
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run('Farklılaştırma ve Kapsayıcılık')
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = self.text_color

            # Turkish field name mappings for farklılaştırma
            fark_field_map = {
                'ogrenme_hizi_farki': 'Öğrenme Hızı Farkı',
                'ileri_duzey': 'İleri Düzey',
                'temel_duzey': 'Temel Düzey',
                'dil_destegi': 'Dil Desteği',
                'kulturel_kapsayicilik': 'Kültürel Kapsayıcılık',
                'sosyal_duygusal_farklilastirma': 'Sosyal Duygusal Farklılaştırma',
                'cekingen_cocuklar': 'Çekingen Çocuklar',
                'dis_donuk_cocuklar': 'Dış Dönük Çocuklar',
                'duygusal_destek': 'Duygusal Destek',
                'coklu_duyusal_yaklasim': 'Çoklu Duyusal Yaklaşım'
            }

            # Parse farklilastirma_kapsayicilik text
            fark_text = json_data.get('farklilastirma_kapsayicilik', '')
            if fark_text and isinstance(fark_text, str):
                # Split and format the text
                sections = fark_text.split('\n\n')
                for section in sections:
                    if section.strip():
                        lines = section.strip().split('\n')
                        for line in lines:
                            line = line.strip()
                            if line.endswith(':'):
                                p = doc.add_paragraph()
                                run = p.add_run(line)
                                run.font.bold = True
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                                run.font.color.rgb = self.text_color
                            elif line.startswith('•'):
                                p = doc.add_paragraph(style='List Bullet')
                                run = p.add_run(line[1:].strip())
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                                run.font.color.rgb = self.text_color
                            elif line:
                                p = doc.add_paragraph()
                                run = p.add_run(line)
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                                run.font.color.rgb = self.text_color

            # Legacy JSON format
            farklilastirma = json_data.get('farklilastirma_ve_kapsayicilik', {})
            if isinstance(farklilastirma, dict):
                for key, value in farklilastirma.items():
                    p = doc.add_paragraph()
                    # If key already has Turkish characters, use it as is
                    if any(char in key for char in 'ÇĞÖŞÜçğöşüıİ'):
                        field_name = key
                    else:
                        field_name = fark_field_map.get(key, key.replace('_', ' ').title())
                    p.add_run(f'{field_name}:').bold = True

                    if isinstance(value, dict):
                        for sub_key, sub_value in value.items():
                            p = doc.add_paragraph(style='List Bullet')
                            # If sub_key already has Turkish characters, use it as is
                            if any(char in sub_key for char in 'ÇĞÖŞÜçğöşüıİ'):
                                sub_field_name = sub_key
                            else:
                                sub_field_name = fark_field_map.get(sub_key, sub_key.replace('_', ' ').title())
                            p.add_run(f'{sub_field_name}: ')
                            p.add_run(sub_value)
                    else:
                        p = doc.add_paragraph()
                        p.add_run(value)
            
            # Evaluation
            self.add_styled_header(doc, 'DEĞERLENDİRME', page_break=True)

            # Parse evaluation text fields
            eval_fields = [
                ('degerlendirme_program', 'Program Değerlendirmesi'),
                ('degerlendirme_beceriler', 'Beceri Değerlendirmesi'),
                ('degerlendirme_ogrenciler', 'Öğrenci Değerlendirmesi')
            ]

            for field_key, field_title in eval_fields:
                field_text = json_data.get(field_key, '')
                if field_text:
                    para = doc.add_paragraph()
                    run = para.add_run(field_title + ':')
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = self.text_color

                    # Parse and add content
                    lines = field_text.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line.endswith(':'):
                            p = doc.add_paragraph()
                            run = p.add_run(line)
                            run.font.bold = True
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                            run.font.color.rgb = self.text_color
                        elif line.startswith('•'):
                            p = doc.add_paragraph(style='List Bullet')
                            run = p.add_run(line[1:].strip())
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                            run.font.color.rgb = self.text_color
                        elif line:
                            p = doc.add_paragraph()
                            run = p.add_run(line)
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                            run.font.color.rgb = self.text_color

                    doc.add_paragraph()  # Add spacing between sections

            # Legacy JSON format
            degerlendirme = json_data.get('degerlendirme', {})
            if isinstance(degerlendirme, dict):
                # Support both old and new formats for program evaluation
                if degerlendirme.get('gozlem_formu'):
                    p = doc.add_paragraph()
                    p.add_run('Gözlem Formu:').bold = True
                    for item in degerlendirme['gozlem_formu']:
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(item)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.color.rgb = self.text_color
                elif degerlendirme.get('program_tarafindan'):
                    p = doc.add_paragraph()
                    p.add_run('Program Değerlendirmesi:').bold = True
                    for item in degerlendirme['program_tarafindan']:
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(item)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.color.rgb = self.text_color
                elif activity_data.get('degerlendirme_program'):
                    p = doc.add_paragraph()
                    p.add_run('Program Değerlendirmesi:').bold = True
                    p = doc.add_paragraph()
                    p.add_run(activity_data['degerlendirme_program'])

                # Support both old and new formats for skills evaluation
                if degerlendirme.get('ogretmen_oz_degerlendirme'):
                    p = doc.add_paragraph()
                    p.add_run('Öğretmen Öz Değerlendirme:').bold = True
                    for item in degerlendirme['ogretmen_oz_degerlendirme']:
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(item)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.color.rgb = self.text_color
                elif degerlendirme.get('amaclanan_beceriler_tarafindan'):
                    p = doc.add_paragraph()
                    p.add_run('Beceri Değerlendirmesi:').bold = True
                    for item in degerlendirme['amaclanan_beceriler_tarafindan']:
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(item)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.color.rgb = self.text_color
                elif activity_data.get('degerlendirme_beceriler'):
                    p = doc.add_paragraph()
                    p.add_run('Beceriler Değerlendirmesi:').bold = True
                    p = doc.add_paragraph()
                    p.add_run(activity_data['degerlendirme_beceriler'])

                # Support both old and new formats for student evaluation
                if degerlendirme.get('cocuk_degerlendirmesi'):
                    p = doc.add_paragraph()
                    p.add_run('Çocuk Değerlendirmesi:').bold = True
                    for item in degerlendirme['cocuk_degerlendirmesi']:
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(item)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.color.rgb = self.text_color
                elif degerlendirme.get('ogrenciler_tarafindan'):
                    p = doc.add_paragraph()
                    p.add_run('Öğrenci Değerlendirmesi:').bold = True
                    for item in degerlendirme['ogrenciler_tarafindan']:
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(item)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.color.rgb = self.text_color
                elif activity_data.get('degerlendirme_ogrenciler'):
                    p = doc.add_paragraph()
                    p.add_run('Öğrenci Değerlendirmesi:').bold = True
                    p = doc.add_paragraph()
                    p.add_run(activity_data['degerlendirme_ogrenciler'])

                # Add family feedback if available
                if degerlendirme.get('aile_geri_bildirimi'):
                    p = doc.add_paragraph()
                    p.add_run('Aile Geri Bildirimi:').bold = True
                    for item in degerlendirme['aile_geri_bildirimi']:
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(item)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.color.rgb = self.text_color
            
            # Save to BytesIO
            docx_file = BytesIO()
            doc.save(docx_file)
            docx_file.seek(0)
            
            return docx_file
            
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {str(e)}")
            raise

# Singleton instance
docx_export_service = DocxExportService()