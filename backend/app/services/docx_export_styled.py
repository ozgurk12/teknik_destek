# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import json
from typing import Dict, Optional
import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)

class StyledDocxExportService:
    """Service for exporting activities to styled DOCX format matching EduPage Kids template"""

    def __init__(self):
        # Color scheme from template
        self.primary_color = RGBColor(114, 191, 190)  # Turquoise/Teal
        self.text_color = RGBColor(44, 62, 80)  # Dark blue-gray
        self.header_bg_color = RGBColor(162, 216, 216)  # Light turquoise

    def add_logo(self, doc, header_paragraph):
        """Add EduPage Kids logo to header"""
        try:
            # Try to add logo if it exists
            logo_path = Path(__file__).parent.parent.parent / "assets" / "edupage_kids_logo.png"
            if logo_path.exists():
                run = header_paragraph.add_run()
                run.add_picture(str(logo_path), width=Inches(1.5))
        except Exception as e:
            # If logo doesn't exist, add text version
            run = header_paragraph.add_run("EduPage\nKIDS")
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = self.primary_color

    def set_cell_background(self, cell, color_hex):
        """Set background color for a table cell"""
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), color_hex)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def create_header_table(self, doc, title):
        """Create styled header with title and logo"""
        # Create a table for the header
        header_table = doc.add_table(rows=1, cols=2)
        header_table.style = 'Table Grid'
        header_table.allow_autofit = False

        # Set column widths
        header_table.columns[0].width = Cm(14)
        header_table.columns[1].width = Cm(4)

        # Title cell
        title_cell = header_table.cell(0, 0)
        self.set_cell_background(title_cell, "B2D8D8")
        title_para = title_cell.paragraphs[0]
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_para.add_run(title.upper())
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = self.text_color

        # Logo cell
        logo_cell = header_table.cell(0, 1)
        self.set_cell_background(logo_cell, "B2D8D8")
        logo_para = logo_cell.paragraphs[0]
        logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.add_logo(doc, logo_para)

        # Remove borders
        for row in header_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'

        doc.add_paragraph()  # Add spacing

    def add_section_header(self, doc, title):
        """Add a section header with background color"""
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = para.add_run(title)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = self.text_color
        doc.add_paragraph()  # Add spacing

    def add_info_section(self, doc, data):
        """Add basic information section"""
        info_items = [
            ("Alan Adı", data.get('alan_adi', '')),
            ("Yaş Grubu", data.get('yas_grubu', '')),
            ("Süre", f"{data.get('sure', '30')} dakika"),
            ("Uygulama Yeri", data.get('uygulama_yeri', ''))
        ]

        for label, value in info_items:
            para = doc.add_paragraph()
            # Label
            label_run = para.add_run(f"{label}: ")
            label_run.font.name = 'Arial'
            label_run.font.size = Pt(12)
            label_run.font.bold = True
            label_run.font.color.rgb = self.text_color
            # Value
            value_run = para.add_run(value)
            value_run.font.name = 'Arial'
            value_run.font.size = Pt(12)
            value_run.font.color.rgb = self.text_color

        doc.add_paragraph()  # Add spacing

    def add_objectives_section(self, doc, data):
        """Add objectives section"""
        self.add_section_header(doc, "Etkinliğin Amacı:")

        amaclar = data.get('amaclar', [])
        for amac in amaclar:
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(amac)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.color.rgb = self.text_color

    def add_materials_section(self, doc, data):
        """Add materials section"""
        self.add_section_header(doc, "Materyaller")

        materyaller = data.get('materyaller', [])
        if isinstance(materyaller, list):
            for item in materyaller:
                para = doc.add_paragraph(style='List Bullet')

                # Check if item has alternatives
                if "Alternatif:" in item:
                    parts = item.split("Alternatif:")
                    main_run = para.add_run(parts[0].strip())
                    main_run.font.name = 'Arial'
                    main_run.font.size = Pt(11)
                    main_run.font.color.rgb = self.text_color

                    if len(parts) > 1:
                        alt_run = para.add_run("\n    Alternatif: " + parts[1].strip())
                        alt_run.font.name = 'Arial'
                        alt_run.font.size = Pt(10)
                        alt_run.font.italic = True
                        alt_run.font.color.rgb = self.text_color
                else:
                    run = para.add_run(item)
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.color.rgb = self.text_color

    def add_implementation_process(self, doc, data):
        """Add implementation process section"""
        doc.add_page_break()
        self.create_header_table(doc, data.get('etkinlik_adi', 'ETKİNLİK'))

        self.add_section_header(doc, "Uygulama Süreci")

        surec = data.get('uygulama_sureci', {})

        # Giriş
        if surec.get('giris'):
            giris = surec['giris']
            para = doc.add_paragraph()
            title_run = para.add_run(f"Giriş ({giris.get('sure', '5 dk')})")
            title_run.font.name = 'Arial'
            title_run.font.size = Pt(12)
            title_run.font.bold = True
            title_run.font.color.rgb = self.text_color

            for adim in giris.get('adimlar', []):
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(adim)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = self.text_color
            doc.add_paragraph()

        # Gelişme
        if surec.get('gelisme'):
            gelisme = surec['gelisme']
            para = doc.add_paragraph()
            title_run = para.add_run(f"Gelişme ({gelisme.get('sure', '15 dk')})")
            title_run.font.name = 'Arial'
            title_run.font.size = Pt(12)
            title_run.font.bold = True
            title_run.font.color.rgb = self.text_color

            aktiviteler = gelisme.get('aktiviteler', {})
            counter = 1
            for key, value in aktiviteler.items():
                # Add numbered subsection
                para = doc.add_paragraph()
                number_run = para.add_run(f"{counter}. ")
                number_run.font.name = 'Arial'
                number_run.font.size = Pt(11)
                number_run.font.bold = True
                number_run.font.color.rgb = self.text_color

                # Extract title and content
                if ":" in value:
                    parts = value.split(":", 1)
                    title = parts[0].replace("_", " ").title()
                    content = parts[1].strip() if len(parts) > 1 else ""

                    title_run = para.add_run(f"{title}")
                    title_run.font.name = 'Arial'
                    title_run.font.size = Pt(11)
                    title_run.font.bold = True
                    title_run.font.color.rgb = self.text_color

                    if content:
                        content_para = doc.add_paragraph()
                        content_para.paragraph_format.left_indent = Pt(18)
                        content_run = content_para.add_run(content)
                        content_run.font.name = 'Arial'
                        content_run.font.size = Pt(11)
                        content_run.font.color.rgb = self.text_color
                else:
                    content_run = para.add_run(value)
                    content_run.font.name = 'Arial'
                    content_run.font.size = Pt(11)
                    content_run.font.color.rgb = self.text_color

                counter += 1
            doc.add_paragraph()

        # Yansıma Çemberi
        if surec.get('yansima_cemberi'):
            yansima = surec['yansima_cemberi']
            para = doc.add_paragraph()
            title_run = para.add_run(f"Yansıma Çemberi ({yansima.get('sure', '5 dk')})")
            title_run.font.name = 'Arial'
            title_run.font.size = Pt(12)
            title_run.font.bold = True
            title_run.font.color.rgb = self.text_color

            if yansima.get('duzenleme'):
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(yansima['duzenleme'])
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = self.text_color

            if yansima.get('sorular'):
                para = doc.add_paragraph()
                run = para.add_run("Öğretmen şu soruları sorar:")
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = self.text_color

                for soru in yansima['sorular']:
                    para = doc.add_paragraph(style='List Bullet')
                    run = para.add_run(f'"{soru}"')
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.color.rgb = self.text_color
            doc.add_paragraph()

        # Sonuç
        if surec.get('sonuc'):
            sonuc = surec['sonuc']
            para = doc.add_paragraph()
            title_run = para.add_run(f"Sonuç ({sonuc.get('sure', '5 dk')})")
            title_run.font.name = 'Arial'
            title_run.font.size = Pt(12)
            title_run.font.bold = True
            title_run.font.color.rgb = self.text_color

            for aktivite in sonuc.get('aktiviteler', []):
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(aktivite)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = self.text_color

    def add_differentiation_section(self, doc, data):
        """Add differentiation and inclusion section"""
        doc.add_page_break()
        self.create_header_table(doc, "FARKLILAŞTIRMA VE KAPSAYICILIK")

        farklilastirma = data.get('farklilastirma_ve_kapsayicilik', {})

        for key, value in farklilastirma.items():
            # Format the key
            formatted_key = key.replace('_', ' ').title()
            if 'ogrenme_hizi' in key:
                formatted_key = "Öğrenme Hızı Farkı:"
            elif 'dil_destegi' in key:
                formatted_key = "Dil Desteği:"
            elif 'kulturel' in key:
                formatted_key = "Kültürel Kapsayıcılık:"
            elif 'sosyal_duygusal' in key:
                formatted_key = "Sosyal-Duygusal Farklılaştırma:"
            elif 'coklu_duyusal' in key:
                formatted_key = "Çoklu Duyusal Yaklaşım:"

            para = doc.add_paragraph()
            title_run = para.add_run(formatted_key)
            title_run.font.name = 'Arial'
            title_run.font.size = Pt(12)
            title_run.font.bold = True
            title_run.font.color.rgb = self.text_color

            if isinstance(value, dict):
                for sub_key, sub_value in value.items():
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Pt(18)

                    sub_formatted = sub_key.replace('_', ' ').title()
                    if 'ileri' in sub_key:
                        sub_formatted = "İleri düzey"
                    elif 'temel' in sub_key:
                        sub_formatted = "Temel düzey"
                    elif 'cekingen' in sub_key:
                        sub_formatted = "Çekingen çocuklar"
                    elif 'dis_donuk' in sub_key:
                        sub_formatted = "Dış dönük çocuklar"

                    sub_run = para.add_run(f"{sub_formatted}: ")
                    sub_run.font.name = 'Arial'
                    sub_run.font.size = Pt(11)
                    sub_run.font.bold = True
                    sub_run.font.color.rgb = self.text_color

                    content_run = para.add_run(sub_value)
                    content_run.font.name = 'Arial'
                    content_run.font.size = Pt(11)
                    content_run.font.color.rgb = self.text_color
            else:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Pt(18)
                run = para.add_run(value)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = self.text_color

            doc.add_paragraph()  # Add spacing

    def add_evaluation_section(self, doc, data):
        """Add evaluation section"""
        doc.add_page_break()
        self.create_header_table(doc, "DEĞERLENDİRME")

        degerlendirme = data.get('degerlendirme', {})

        # Program Tarafından
        if degerlendirme.get('gozlem_formu') or degerlendirme.get('program_tarafindan'):
            para = doc.add_paragraph()
            title_run = para.add_run("Program Tarafından:")
            title_run.font.name = 'Arial'
            title_run.font.size = Pt(12)
            title_run.font.bold = True
            title_run.font.color.rgb = self.text_color

            items = degerlendirme.get('gozlem_formu', degerlendirme.get('program_tarafindan', []))
            for item in items:
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(item)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = self.text_color
            doc.add_paragraph()

        # Amaçlanan Beceriler Tarafından
        if degerlendirme.get('ogretmen_oz_degerlendirme') or degerlendirme.get('amaclanan_beceriler_tarafindan'):
            para = doc.add_paragraph()
            title_run = para.add_run("Amaçlanan Beceriler Tarafından:")
            title_run.font.name = 'Arial'
            title_run.font.size = Pt(12)
            title_run.font.bold = True
            title_run.font.color.rgb = self.text_color

            items = degerlendirme.get('ogretmen_oz_degerlendirme', degerlendirme.get('amaclanan_beceriler_tarafindan', []))
            for item in items:
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(item)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = self.text_color
            doc.add_paragraph()

        # Öğrenciler Tarafından
        if degerlendirme.get('cocuk_degerlendirmesi') or degerlendirme.get('ogrenciler_tarafindan'):
            para = doc.add_paragraph()
            title_run = para.add_run("Öğrenciler Tarafından:")
            title_run.font.name = 'Arial'
            title_run.font.size = Pt(12)
            title_run.font.bold = True
            title_run.font.color.rgb = self.text_color

            items = degerlendirme.get('cocuk_degerlendirmesi', degerlendirme.get('ogrenciler_tarafindan', []))
            for item in items:
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(item)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = self.text_color

    def export_activity_to_docx(self, activity_data: Dict) -> BytesIO:
        """Export activity data to styled DOCX format"""
        try:
            # Parse JSON data if it's stored as string
            if isinstance(activity_data.get('json_data'), str):
                json_data = json.loads(activity_data['json_data'])
            else:
                json_data = activity_data

            # Create a new Document
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)

            # Add header with title and logo
            self.create_header_table(doc, json_data.get('etkinlik_adi', 'ETKİNLİK'))

            # Add basic information
            self.add_info_section(doc, json_data)

            # Add objectives
            self.add_objectives_section(doc, json_data)

            # Add materials
            self.add_materials_section(doc, json_data)

            # Add implementation process
            self.add_implementation_process(doc, json_data)

            # Add differentiation section
            self.add_differentiation_section(doc, json_data)

            # Add evaluation section
            self.add_evaluation_section(doc, json_data)

            # Save to BytesIO
            docx_file = BytesIO()
            doc.save(docx_file)
            docx_file.seek(0)

            return docx_file

        except Exception as e:
            logger.error(f"Error exporting to styled DOCX: {str(e)}")
            raise

# Singleton instance
styled_docx_export_service = StyledDocxExportService()