# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from io import BytesIO
import json
from typing import Dict, Optional
import logging
from pathlib import Path
import os
from lxml import etree

logger = logging.getLogger(__name__)

class DocxExportService:
    """Service for exporting activities to DOCX format with watermark"""

    def __init__(self):
        # Color scheme from template
        self.primary_color = RGBColor(114, 191, 190)  # Turquoise/Teal
        self.text_color = RGBColor(44, 62, 80)  # Dark blue-gray
        self.background_image_path = '/Users/ozgurk12/Desktop/PROJELER/EduPage Kids/Etkinlik ve Günlük Plan/Varlık 2@3x.png'

    def add_watermark(self, doc):
        """Add watermark image to document background"""
        try:
            if not os.path.exists(self.background_image_path):
                logger.warning(f"Background image not found")
                return

            # Access document's main section
            section = doc.sections[0]

            # Different approach: modify header to contain full-page background
            header = section.header

            # Set header margin to 0
            section.header_distance = Cm(0)

            # Clear header content
            for p in header.paragraphs:
                p._element.getparent().remove(p._element)

            # Add new paragraph with image
            p = header.add_paragraph()
            r = p.add_run()

            # Add picture inline first
            pic = r.add_picture(self.background_image_path, width=Cm(21))

            # Now convert to floating behind text
            inline = pic._inline

            # Create XML for anchor (floating image)
            # namespace map
            nsmap = {
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
            }

            # Build anchor element
            anchor = etree.Element('{%s}anchor' % nsmap['wp'], nsmap=nsmap)
            anchor.set('distT', '0')
            anchor.set('distB', '0')
            anchor.set('distL', '0')
            anchor.set('distR', '0')
            anchor.set('simplePos', '0')
            anchor.set('relativeHeight', '0')
            anchor.set('behindDoc', '1')  # This puts it behind text
            anchor.set('locked', '0')
            anchor.set('layoutInCell', '1')
            anchor.set('allowOverlap', '1')

            # Position
            simplePos = etree.SubElement(anchor, '{%s}simplePos' % nsmap['wp'])
            simplePos.set('x', '0')
            simplePos.set('y', '0')

            posH = etree.SubElement(anchor, '{%s}positionH' % nsmap['wp'])
            posH.set('relativeFrom', 'page')
            posOffset = etree.SubElement(posH, '{%s}posOffset' % nsmap['wp'])
            posOffset.text = '0'

            posV = etree.SubElement(anchor, '{%s}positionV' % nsmap['wp'])
            posV.set('relativeFrom', 'page')
            posOffsetV = etree.SubElement(posV, '{%s}posOffset' % nsmap['wp'])
            posOffsetV.text = '0'

            # Extent (size)
            extent = etree.SubElement(anchor, '{%s}extent' % nsmap['wp'])
            extent.set('cx', str(inline.extent.cx))
            extent.set('cy', str(inline.extent.cy))

            # Effect extent
            effectExtent = etree.SubElement(anchor, '{%s}effectExtent' % nsmap['wp'])
            effectExtent.set('l', '0')
            effectExtent.set('t', '0')
            effectExtent.set('r', '0')
            effectExtent.set('b', '0')

            # Wrap none
            wrapNone = etree.SubElement(anchor, '{%s}wrapNone' % nsmap['wp'])

            # DocPr
            docPr = etree.SubElement(anchor, '{%s}docPr' % nsmap['wp'])
            docPr.set('id', str(inline.docPr.id))
            docPr.set('name', inline.docPr.name)

            # Graphic frame properties
            cNvGraphicFramePr = etree.SubElement(anchor, '{%s}cNvGraphicFramePr' % nsmap['wp'])
            graphicFrameLocks = etree.SubElement(cNvGraphicFramePr, '{%s}graphicFrameLocks' % nsmap['a'])
            graphicFrameLocks.set('noChangeAspect', '1')

            # Move graphic from inline to anchor
            graphic = inline.graphic._element
            anchor.append(graphic)

            # Replace inline with anchor
            r._element.replace(inline._element, anchor)

        except Exception as e:
            logger.error(f"Error adding watermark: {e}")

    def parse_text_to_list(self, text: str) -> list:
        """Parse bullet-pointed text into a list"""
        if not text:
            return []

        lines = text.split('\n')
        result = []
        for line in lines:
            line = line.strip()
            if line.startswith('•'):
                result.append(line[1:].strip())
            elif line:
                result.append(line)
        return result

    def export_activity_to_docx(self, activity_data: Dict) -> BytesIO:
        """Export activity data to DOCX format with watermark background"""
        try:
            # Use activity_data directly
            json_data = activity_data

            # Create document
            doc = Document()

            # Set margins
            for section in doc.sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)

            # Add watermark
            self.add_watermark(doc)

            # Set default font
            doc.styles['Normal'].font.name = 'Arial'
            doc.styles['Normal'].font.color.rgb = self.text_color

            # Title
            doc.add_paragraph()  # Spacing
            doc.add_paragraph()  # Spacing

            title_para = doc.add_paragraph()
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title_para.add_run(json_data.get('etkinlik_adi', 'ETKİNLİK').upper())
            title_run.font.name = 'Arial'
            title_run.font.size = Pt(24)
            title_run.font.bold = True
            title_run.font.color.rgb = self.primary_color

            doc.add_paragraph()  # Spacing

            # Basic Information Table
            para = doc.add_paragraph()
            run = para.add_run('Temel Bilgiler')
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = self.text_color

            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Table Grid'

            info_data = [
                ('Alan Adı:', json_data.get('alan_adi', '')),
                ('Yaş Grubu:', json_data.get('yas_grubu', '')),
                ('Süre:', f"{json_data.get('sure', '30')} dakika"),
                ('Uygulama Yeri:', json_data.get('uygulama_yeri', ''))
            ]

            for i, (label, value) in enumerate(info_data):
                row = info_table.rows[i]
                # Label
                cell_para = row.cells[0].paragraphs[0]
                cell_run = cell_para.add_run(label)
                cell_run.font.name = 'Arial'
                cell_run.font.size = Pt(11)
                cell_run.font.bold = True
                # Value
                cell_para = row.cells[1].paragraphs[0]
                cell_run = cell_para.add_run(value)
                cell_run.font.name = 'Arial'
                cell_run.font.size = Pt(11)

            # Objectives
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run('Etkinliğin Amacı:')
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True

            amaclar = self.parse_text_to_list(json_data.get('etkinlik_amaci', ''))
            for amac in amaclar:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(amac)
                run.font.name = 'Arial'
                run.font.size = Pt(11)

            # Materials
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run('Materyaller')
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True

            materyaller = self.parse_text_to_list(json_data.get('materyaller', ''))
            for item in materyaller:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(item)
                run.font.name = 'Arial'
                run.font.size = Pt(11)

            # Implementation Process
            doc.add_page_break()
            self.add_watermark(doc)  # Add watermark to new page

            doc.add_paragraph()
            heading = doc.add_heading('UYGULAMA SÜRECİ', level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in heading.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(20)
                run.font.color.rgb = self.primary_color

            surec_text = json_data.get('uygulama_sureci', '')
            if surec_text:
                sections = surec_text.split('\n\n')
                for section in sections:
                    if section.strip():
                        lines = section.strip().split('\n')
                        if lines:
                            # Section header
                            header = lines[0]
                            para = doc.add_paragraph()
                            run = para.add_run(header)
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                            run.font.bold = True

                            # Content
                            for line in lines[1:]:
                                line = line.strip()
                                if line.startswith('•'):
                                    p = doc.add_paragraph(style='List Bullet')
                                    run = p.add_run(line[1:].strip())
                                elif line:
                                    p = doc.add_paragraph()
                                    run = p.add_run(line)
                                if line:
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)

            # Adaptations
            if json_data.get('uyarlama'):
                doc.add_page_break()
                self.add_watermark(doc)

                heading = doc.add_heading('UYARLAMA VE FARKLILAŞTIRMA', level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in heading.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(20)
                    run.font.color.rgb = self.primary_color

                uyarlama_text = json_data.get('uyarlama', '')
                if uyarlama_text:
                    sections = uyarlama_text.split('\n\n')
                    for section in sections:
                        if section.strip():
                            lines = section.strip().split('\n')
                            for line in lines:
                                line = line.strip()
                                if line.endswith(':'):
                                    p = doc.add_paragraph()
                                    run = p.add_run(line)
                                    run.font.bold = True
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                elif line.startswith('•'):
                                    p = doc.add_paragraph(style='List Bullet')
                                    run = p.add_run(line[1:].strip())
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                elif line:
                                    p = doc.add_paragraph()
                                    run = p.add_run(line)
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)

                # Farklılaştırma
                fark_text = json_data.get('farklilastirma_kapsayicilik', '')
                if fark_text:
                    doc.add_paragraph()
                    para = doc.add_paragraph()
                    run = para.add_run('Farklılaştırma ve Kapsayıcılık')
                    run.font.name = 'Arial'
                    run.font.size = Pt(14)
                    run.font.bold = True

                    sections = fark_text.split('\n\n')
                    for section in sections:
                        if section.strip():
                            lines = section.strip().split('\n')
                            for line in lines:
                                line = line.strip()
                                if line.endswith(':'):
                                    p = doc.add_paragraph()
                                    run = p.add_run(line)
                                    run.font.bold = True
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                elif line.startswith('•'):
                                    p = doc.add_paragraph(style='List Bullet')
                                    run = p.add_run(line[1:].strip())
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)
                                elif line:
                                    p = doc.add_paragraph()
                                    run = p.add_run(line)
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(11)

            # Evaluation
            if any([json_data.get('degerlendirme_program'),
                    json_data.get('degerlendirme_beceriler'),
                    json_data.get('degerlendirme_ogrenciler')]):
                doc.add_page_break()
                self.add_watermark(doc)

                heading = doc.add_heading('DEĞERLENDİRME', level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in heading.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(20)
                    run.font.color.rgb = self.primary_color

                eval_fields = [
                    ('degerlendirme_program', 'Program Değerlendirmesi'),
                    ('degerlendirme_beceriler', 'Beceri Değerlendirmesi'),
                    ('degerlendirme_ogrenciler', 'Öğrenci Değerlendirmesi')
                ]

                for field_key, field_title in eval_fields:
                    field_text = json_data.get(field_key, '')
                    if field_text:
                        para = doc.add_paragraph()
                        run = para.add_run(field_title + ':')
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
                        run.font.bold = True

                        lines = field_text.split('\n')
                        for line in lines:
                            line = line.strip()
                            if line.endswith(':'):
                                p = doc.add_paragraph()
                                run = p.add_run(line)
                                run.font.bold = True
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                            elif line.startswith('•'):
                                p = doc.add_paragraph(style='List Bullet')
                                run = p.add_run(line[1:].strip())
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                            elif line:
                                p = doc.add_paragraph()
                                run = p.add_run(line)
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)

                        doc.add_paragraph()

            # Save to BytesIO
            docx_file = BytesIO()
            doc.save(docx_file)
            docx_file.seek(0)

            return docx_file

        except Exception as e:
            logger.error(f"Error exporting to DOCX: {str(e)}")
            raise

# Singleton instance
docx_export_service = DocxExportService()