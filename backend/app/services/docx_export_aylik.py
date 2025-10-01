import os
import tempfile
from pathlib import Path
from typing import Any, Dict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import logging

logger = logging.getLogger(__name__)

class AylikPlanDocxExporter:
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()

    def export_aylik_plan(self, plan: Any) -> str:
        """
        Aylık planı DOCX formatına dönüştürür.
        """
        try:
            # Yeni bir dokuman oluştur
            doc = Document()

            # Başlık ekle
            self._add_title(doc, plan.plan_adi)

            # Plan bilgileri
            self._add_plan_info(doc, plan)

            # Alan Becerileri
            self._add_section(doc, "1. ALAN BECERİLERİ", plan.alan_becerileri)

            # Kavramsal Beceriler
            self._add_section(doc, "2. KAVRAMSAL BECERİLER", plan.kavramsal_beceriler)

            # Eğilimler
            self._add_section(doc, "3. EĞİLİMLER", plan.egilimler)

            # Sosyal-Duygusal Öğrenme Becerileri
            self._add_section(doc, "4. SOSYAL-DUYGUSAL ÖĞRENME BECERİLERİ", plan.sosyal_duygusal_beceriler)

            # Değerler
            self._add_section(doc, "5. DEĞERLER", plan.degerler)

            # Okuryazarlık Becerileri
            self._add_section(doc, "6. OKURYAZARLIK BECERİLERİ", plan.okuryazarlik_becerileri)

            # Öğrenme Çıktıları
            self._add_section(doc, "7. ÖĞRENME ÇIKTILARI VE SÜREÇ BİLEŞENLERİ", plan.ogrenme_ciktilari)

            # Anahtar Kavramlar
            self._add_keywords(doc, "8. ANAHTAR KAVRAMLAR", plan.anahtar_kavramlar)

            # Değerlendirme
            self._add_section(doc, "9. ÖĞRENME KANITLARI (ÖLÇME VE DEĞERLENDİRME)", plan.degerlendirme)

            # Öğrenme-Öğretme Yaşantıları
            self._add_text_section(doc, "10. ÖĞRENME-ÖĞRETME YAŞANTILARI", plan.ogrenme_ogretme_yasantilari)

            # Farklılaştırma ve Zenginleştirme
            if plan.farklilastirma_zenginlestirme:
                self._add_text_section(doc, "11. FARKLILAŞTIRMA VE ZENGİNLEŞTİRME", plan.farklilastirma_zenginlestirme)

            # Destekleme
            if plan.destekleme:
                self._add_text_section(doc, "12. DESTEKLEME", plan.destekleme)

            # Aile/Toplum Katılımı
            if plan.aile_toplum_katilimi:
                self._add_text_section(doc, "13. AİLE/TOPLUM KATILIMI", plan.aile_toplum_katilimi)

            # Dosyayı kaydet
            file_path = os.path.join(self.temp_dir, f"{plan.id}_aylik_plan.docx")
            doc.save(file_path)

            return file_path

        except Exception as e:
            logger.error(f"DOCX export hatası: {str(e)}")
            raise Exception(f"DOCX export hatası: {str(e)}")

    def _add_title(self, doc: Document, title: str):
        """Başlık ekler"""
        title_paragraph = doc.add_heading(title, level=0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    def _add_plan_info(self, doc: Document, plan: Any):
        """Plan bilgilerini ekler"""
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'

        # Yaş Grubu
        row = info_table.rows[0]
        row.cells[0].text = "Yaş Grubu:"
        row.cells[1].text = plan.yas_grubu

        # Ay
        row = info_table.rows[1]
        row.cells[0].text = "Ay:"
        row.cells[1].text = plan.ay

        # Yıl
        row = info_table.rows[2]
        row.cells[0].text = "Yıl:"
        row.cells[1].text = str(plan.yil)

        doc.add_paragraph()

    def _add_section(self, doc: Document, title: str, content: Dict):
        """Bölüm ekler"""
        doc.add_heading(title, level=1)

        if isinstance(content, dict):
            for key, value in content.items():
                # Alt başlık
                doc.add_heading(key, level=2)

                if isinstance(value, list):
                    for item in value:
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(str(item))
                elif isinstance(value, str):
                    doc.add_paragraph(value)
                else:
                    doc.add_paragraph(str(value))
        else:
            doc.add_paragraph(str(content))

        doc.add_paragraph()

    def _add_keywords(self, doc: Document, title: str, keywords: list):
        """Anahtar kavramları ekler"""
        doc.add_heading(title, level=1)

        if keywords:
            keywords_text = ", ".join(keywords)
            doc.add_paragraph(keywords_text)
        else:
            doc.add_paragraph("Belirtilmemiş")

        doc.add_paragraph()

    def _add_text_section(self, doc: Document, title: str, content: str):
        """Metin bölümü ekler"""
        doc.add_heading(title, level=1)

        if content:
            # Satırları böl ve her birini ayrı paragraf olarak ekle
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph("Belirtilmemiş")

        doc.add_paragraph()

    def cleanup(self):
        """Geçici dosyaları temizle"""
        import shutil
        try:
            shutil.rmtree(self.temp_dir)
        except Exception as e:
            logger.error(f"Temp dizin temizleme hatası: {str(e)}")