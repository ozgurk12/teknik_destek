# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import json
from typing import Dict, Optional, List, Any
import logging
from pathlib import Path
from datetime import date

logger = logging.getLogger(__name__)

class DocxGunlukExportService:
    """Service for exporting daily plans to DOCX format following MEB Maarif Model"""

    def __init__(self):
        # Color scheme
        self.primary_color = RGBColor(0, 102, 204)  # Blue for main headers
        self.secondary_color = RGBColor(114, 191, 190)  # Turquoise/Teal
        self.text_color = RGBColor(44, 62, 80)  # Dark blue-gray
        self.header_bg_color = RGBColor(240, 240, 240)  # Light gray background

    def parse_json_field(self, field: Any) -> Any:
        """Parse JSON field if it's a string"""
        if isinstance(field, str):
            try:
                return json.loads(field)
            except:
                return field
        return field

    def export_gunluk_plan_to_docx(self, plan_data: Dict) -> BytesIO:
        """Export daily plan data to DOCX format following MEB Maarif Model"""
        try:
            doc = Document()

            # Set page margins
            for section in doc.sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
                # A4 size
                section.page_height = Cm(29.7)
                section.page_width = Cm(21)

            # Main Title - MEB MAARİF MODELİ
            title_para = doc.add_paragraph()
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title_para.add_run("MEB MAARİF MODELİ OKUL ÖNCESİ EĞİTİM PROGRAMI")
            title_run.font.name = 'Arial'
            title_run.font.size = Pt(16)
            title_run.font.bold = True
            title_run.font.color.rgb = self.primary_color

            # Subtitle - GÜNLÜK PLAN
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            subtitle_run = subtitle_para.add_run("GÜNLÜK PLAN")
            subtitle_run.font.name = 'Arial'
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.bold = True
            subtitle_run.font.color.rgb = self.text_color

            doc.add_paragraph()

            # Plan Info Table
            info_table = doc.add_table(rows=2, cols=3)
            info_table.style = 'Table Grid'

            # Row 1
            cell = info_table.cell(0, 0)
            cell.text = f"Plan Adı: {plan_data.get('plan_adi', '')}"
            cell = info_table.cell(0, 1)
            if plan_data.get('tarih'):
                if isinstance(plan_data['tarih'], date):
                    date_str = plan_data['tarih'].strftime('%d.%m.%Y')
                else:
                    date_str = str(plan_data['tarih'])
                cell.text = f"Tarih: {date_str}"
            cell = info_table.cell(0, 2)
            cell.text = f"Yaş Grubu: {plan_data.get('yas_grubu', '')}"

            # Row 2
            cell = info_table.cell(1, 0)
            cell.text = f"Sınıf: {plan_data.get('sinif', '')}"
            cell = info_table.cell(1, 1)
            cell.text = f"Öğretmen: {plan_data.get('ogretmen', '')}"
            cell = info_table.cell(1, 2)
            cell.text = ""

            doc.add_paragraph()

            # 1. MÜFREDAT BİLEŞENLERİ
            self._add_main_section_header(doc, "1. MÜFREDAT BİLEŞENLERİ")

            # Alan Becerileri
            alan_becerileri = self.parse_json_field(plan_data.get('alan_becerileri'))
            if alan_becerileri:
                self._add_subsection_header(doc, "Alan Becerileri")
                self._add_curriculum_dict(doc, alan_becerileri)

            # Kavramsal Beceriler
            kavramsal_beceriler = self.parse_json_field(plan_data.get('kavramsal_beceriler'))
            if kavramsal_beceriler:
                self._add_subsection_header(doc, "Kavramsal Beceriler")
                self._add_curriculum_dict(doc, kavramsal_beceriler)

            # Eğilimler
            egilimler = self.parse_json_field(plan_data.get('egilimler'))
            if egilimler:
                self._add_subsection_header(doc, "Eğilimler")
                self._add_curriculum_dict(doc, egilimler)

            doc.add_paragraph()

            # 2. PROGRAMLAR ARASI BİLEŞENLER
            self._add_main_section_header(doc, "2. PROGRAMLAR ARASI BİLEŞENLER")

            programlar_arasi = self.parse_json_field(plan_data.get('programlar_arasi_bilesenler'))
            if programlar_arasi and isinstance(programlar_arasi, dict):
                # Check for lowercase keys (from AI response)
                sosyal = programlar_arasi.get('sosyal_duygusal_ogrenme_becerileri') or programlar_arasi.get('Sosyal Duygusal Beceriler')
                degerler = programlar_arasi.get('degerler') or programlar_arasi.get('Değerler')
                okuryazarlik = programlar_arasi.get('okuryazarlik_becerileri') or programlar_arasi.get('Okuryazarlık Becerileri')

                if sosyal:
                    self._add_subsection_header(doc, "Sosyal Duygusal Öğrenme Becerileri")
                    self._add_curriculum_dict(doc, sosyal)

                if degerler:
                    self._add_subsection_header(doc, "Değerler")
                    self._add_curriculum_dict(doc, degerler)

                if okuryazarlik:
                    self._add_subsection_header(doc, "Okuryazarlık Becerileri")
                    self._add_curriculum_dict(doc, okuryazarlik)
            else:
                # Fallback to individual fields
                sosyal = self.parse_json_field(plan_data.get('sosyal_duygusal_beceriler'))
                degerler_field = self.parse_json_field(plan_data.get('degerler'))
                okuryazarlik_field = self.parse_json_field(plan_data.get('okuryazarlik_becerileri'))

                if sosyal:
                    self._add_subsection_header(doc, "Sosyal Duygusal Öğrenme Becerileri")
                    self._add_curriculum_dict(doc, sosyal)

                if degerler_field:
                    self._add_subsection_header(doc, "Değerler")
                    self._add_curriculum_dict(doc, degerler_field)

                if okuryazarlik_field:
                    self._add_subsection_header(doc, "Okuryazarlık Becerileri")
                    self._add_curriculum_dict(doc, okuryazarlik_field)

                doc.add_paragraph()

            # 3. ÖĞRENME ÇIKTILARI VE SÜREÇ BİLEŞENLERİ
            ogrenme_ciktilari = self.parse_json_field(
                plan_data.get('ogrenme_ciktilari_surec_bilesenleri') or plan_data.get('ogrenme_ciktilari')
            )
            if ogrenme_ciktilari:
                self._add_main_section_header(doc, "3. ÖĞRENME ÇIKTILARI VE SÜREÇ BİLEŞENLERİ")
                self._add_curriculum_dict(doc, ogrenme_ciktilari)
                doc.add_paragraph()

            # 4. İÇERİK ÇERÇEVESİ
            icerik_cercevesi = self.parse_json_field(plan_data.get('icerik_cercevesi'))
            if icerik_cercevesi:
                self._add_main_section_header(doc, "4. İÇERİK ÇERÇEVESİ")

                if isinstance(icerik_cercevesi, dict):
                    if icerik_cercevesi.get('kavramlar'):
                        self._add_content_item(doc, "Kavramlar", icerik_cercevesi['kavramlar'])
                    if icerik_cercevesi.get('sozcukler'):
                        self._add_content_item(doc, "Sözcükler", icerik_cercevesi['sozcukler'])
                    if icerik_cercevesi.get('semboller'):
                        self._add_content_item(doc, "Semboller", icerik_cercevesi['semboller'])
                    if icerik_cercevesi.get('materyaller'):
                        self._add_content_item(doc, "Materyaller", icerik_cercevesi['materyaller'])
                    if icerik_cercevesi.get('egitim_ortamlari'):
                        self._add_content_item(doc, "Eğitim/Öğrenme Ortamları", icerik_cercevesi['egitim_ortamlari'])
            else:
                # Fallback to individual fields
                self._add_main_section_header(doc, "4. İÇERİK ÇERÇEVESİ")
                if plan_data.get('kavramlar'):
                    self._add_content_item(doc, "Kavramlar", plan_data['kavramlar'])
                if plan_data.get('sozcukler'):
                    self._add_content_item(doc, "Sözcükler", plan_data['sozcukler'])
                if plan_data.get('materyaller'):
                    self._add_content_item(doc, "Materyaller", plan_data['materyaller'])
                if plan_data.get('egitim_ortamlari'):
                    self._add_content_item(doc, "Eğitim/Öğrenme Ortamları", plan_data['egitim_ortamlari'])

            doc.add_paragraph()

            # 5. ÖĞRENME-ÖĞRETME YAŞANTILARI
            ogrenme_ogretme = self.parse_json_field(plan_data.get('ogrenme_ogretme_yasantilari'))
            if ogrenme_ogretme:
                self._add_main_section_header(doc, "5. ÖĞRENME-ÖĞRETME YAŞANTILARI")

                if isinstance(ogrenme_ogretme, dict):
                    if ogrenme_ogretme.get('gune_baslama'):
                        self._add_activity_section(doc, "GÜNE BAŞLAMA ZAMANI", ogrenme_ogretme['gune_baslama'])
                    if ogrenme_ogretme.get('ogrenme_merkezleri'):
                        self._add_activity_section(doc, "ÖĞRENME MERKEZLERİNDE OYUN", ogrenme_ogretme['ogrenme_merkezleri'])
                    if ogrenme_ogretme.get('beslenme_toplanma'):
                        self._add_activity_section(doc, "BESLENME, TOPLANMA, TEMİZLİK", ogrenme_ogretme['beslenme_toplanma'])
                    if ogrenme_ogretme.get('etkinlikler'):
                        self._add_activity_section(doc, "ETKİNLİKLER", ogrenme_ogretme['etkinlikler'])
                    if ogrenme_ogretme.get('degerlendirme'):
                        self._add_activity_section(doc, "DEĞERLENDİRME", ogrenme_ogretme['degerlendirme'])
                    if ogrenme_ogretme.get('gunun_degerlendirmesi'):
                        self._add_activity_section(doc, "GÜNÜN DEĞERLENDİRMESİ", ogrenme_ogretme['gunun_degerlendirmesi'])
            else:
                # Fallback to individual fields
                self._add_main_section_header(doc, "5. ÖĞRENME-ÖĞRETME YAŞANTILARI")
                if plan_data.get('gune_baslama'):
                    self._add_activity_section(doc, "GÜNE BAŞLAMA ZAMANI", plan_data['gune_baslama'])
                if plan_data.get('ogrenme_merkezleri'):
                    self._add_activity_section(doc, "ÖĞRENME MERKEZLERİNDE OYUN", plan_data['ogrenme_merkezleri'])
                if plan_data.get('beslenme_toplanma'):
                    self._add_activity_section(doc, "BESLENME, TOPLANMA, TEMİZLİK", plan_data['beslenme_toplanma'])
                if plan_data.get('etkinlikler'):
                    self._add_activity_section(doc, "ETKİNLİKLER", plan_data['etkinlikler'])
                if plan_data.get('degerlendirme'):
                    self._add_activity_section(doc, "DEĞERLENDİRME", plan_data['degerlendirme'])

            doc.add_paragraph()

            # 6. FARKLILASTIRMA
            farklilastirma = self.parse_json_field(plan_data.get('farklilastirma'))
            if farklilastirma and isinstance(farklilastirma, dict):
                self._add_main_section_header(doc, "6. FARKLILASTIRMA")

                zenginlestirme = farklilastirma.get('zenginlestirme')
                destekleme = farklilastirma.get('destekleme')

                if zenginlestirme:
                    self._add_family_section(doc, "Zenginleştirme", zenginlestirme)
                if destekleme:
                    self._add_family_section(doc, "Destekleme", destekleme)

                doc.add_paragraph()
            else:
                # Fallback to individual fields
                if plan_data.get('zenginlestirme') or plan_data.get('destekleme'):
                    self._add_main_section_header(doc, "6. FARKLILASTIRMA")

                    if plan_data.get('zenginlestirme'):
                        self._add_family_section(doc, "Zenginleştirme", plan_data['zenginlestirme'])
                    if plan_data.get('destekleme'):
                        self._add_family_section(doc, "Destekleme", plan_data['destekleme'])

                    doc.add_paragraph()

            # 7. AİLE/TOPLUM KATILIMI
            aile_toplum = self.parse_json_field(plan_data.get('aile_toplum_katilimi'))
            if aile_toplum and isinstance(aile_toplum, dict):
                self._add_main_section_header(doc, "7. AİLE/TOPLUM KATILIMI")

                aile_katilimi = aile_toplum.get('aile_katilimi')
                toplum_katilimi = aile_toplum.get('toplum_katilimi')

                if aile_katilimi:
                    self._add_family_section(doc, "Aile Katılımı", aile_katilimi)
                if toplum_katilimi:
                    self._add_family_section(doc, "Toplum Katılımı", toplum_katilimi)
            else:
                # Fallback to individual fields
                if plan_data.get('aile_katilimi') or plan_data.get('toplum_katilimi'):
                    self._add_main_section_header(doc, "7. AİLE/TOPLUM KATILIMI")

                    if plan_data.get('aile_katilimi'):
                        self._add_family_section(doc, "Aile Katılımı", plan_data['aile_katilimi'])
                    if plan_data.get('toplum_katilimi'):
                        self._add_family_section(doc, "Toplum Katılımı", plan_data['toplum_katilimi'])

            # ÖĞRETMEN NOTLARI
            if plan_data.get('notlar'):
                doc.add_paragraph()
                self._add_main_section_header(doc, "ÖĞRETMEN NOTLARI")
                para = doc.add_paragraph()
                para.add_run(plan_data['notlar']).font.size = Pt(11)

            # Save to BytesIO
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            return output

        except Exception as e:
            logger.error(f"Error creating DOCX: {str(e)}")
            raise

    def _add_main_section_header(self, doc: Document, title: str):
        """Add main section header with formatting"""
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = self.primary_color
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def _add_subsection_header(self, doc: Document, title: str):
        """Add subsection header with formatting"""
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = self.text_color
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def _add_curriculum_dict(self, doc: Document, data: Any):
        """Add curriculum data from dictionary"""
        if isinstance(data, dict):
            for category, items in data.items():
                # Category header
                para = doc.add_paragraph()
                run = para.add_run(f"{category}:")
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.bold = True

                # Items
                if isinstance(items, list):
                    for item in items:
                        item_para = doc.add_paragraph(style='List Bullet')
                        item_para.add_run(str(item)).font.size = Pt(10)
                        item_para.left_indent = Pt(20)
                elif isinstance(items, dict):
                    for sub_cat, sub_items in items.items():
                        sub_para = doc.add_paragraph()
                        sub_para.add_run(f"  {sub_cat}:").font.size = Pt(10)
                        sub_para.left_indent = Pt(20)
                        if isinstance(sub_items, list):
                            for sub_item in sub_items:
                                sub_item_para = doc.add_paragraph(style='List Bullet')
                                sub_item_para.add_run(str(sub_item)).font.size = Pt(10)
                                sub_item_para.left_indent = Pt(40)
                else:
                    item_para = doc.add_paragraph()
                    item_para.add_run(f"  {items}").font.size = Pt(10)
                    item_para.left_indent = Pt(20)
        elif isinstance(data, list):
            for item in data:
                item_para = doc.add_paragraph(style='List Bullet')
                item_para.add_run(str(item)).font.size = Pt(10)

    def _add_content_item(self, doc: Document, label: str, content: str):
        """Add content item with label"""
        para = doc.add_paragraph()
        label_run = para.add_run(f"{label}: ")
        label_run.font.name = 'Arial'
        label_run.font.size = Pt(11)
        label_run.font.bold = True

        content_run = para.add_run(str(content))
        content_run.font.name = 'Arial'
        content_run.font.size = Pt(10)

    def _add_activity_section(self, doc: Document, title: str, content: str):
        """Add activity section with title and content"""
        # Title
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = self.secondary_color

        # Content
        content_para = doc.add_paragraph()
        content_para.add_run(str(content)).font.size = Pt(10)
        content_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        doc.add_paragraph()  # Add spacing

    def _add_family_section(self, doc: Document, title: str, content: str):
        """Add family participation section"""
        para = doc.add_paragraph()
        title_run = para.add_run(f"{title}:")
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(11)
        title_run.font.bold = True

        doc.add_paragraph()
        content_para = doc.add_paragraph()
        content_para.add_run(str(content)).font.size = Pt(10)
        content_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Create singleton instance
docx_gunluk_export_service = DocxGunlukExportService()